from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path("docs/Recovered_45_Day_Execution_Template_Reference.docx")

DETAILED_TASKS = [
    (1, 1, "Site handover, measurement verification, and access confirmation", "Project Setup", "Internal", "Confirm access, working hours, lift rules, storage and society/site permissions."),
    (1, 2, "Vendor kickoff briefing with safety and schedule alignment", "Project Setup", "Internal", "Brief all key vendors on the 45-day plan, safety, access rules and escalation process."),
    (1, 3, "Material storage zone and temporary protection setup", "Civil", "Civil", "Create a safe material storage area and protect finished/working surfaces from damage."),
    (2, 1, "Temporary lighting setup", "Electrical", "Electrical", "Install temporary lighting in primary work zones for safe execution."),
    (2, 2, "Temporary plug points and work power availability", "Electrical", "Electrical", "Confirm safe temporary power points are available to vendors."),
    (2, 3, "Initial floor level and civil condition check", "Civil", "Civil", "Check slab/floor level and note civil preparation requirements."),
    (3, 1, "Demolition or dismantling work start where applicable", "Civil", "Civil", "Start approved dismantling and isolate waste safely."),
    (3, 2, "Debris removal route and housekeeping plan confirmation", "Civil", "Civil", "Confirm debris route, lift timing and housekeeping frequency."),
    (3, 3, "Electrical chase marking and routing confirmation", "Electrical", "Electrical", "Mark electrical routes and confirm with layout."),
    (4, 1, "Civil repair and floor base preparation", "Civil", "Civil", "Prepare civil base and close visible repair points."),
    (4, 2, "Ceiling grid layout marking", "Ceiling", "Ceiling", "Mark ceiling grid based on approved reflected ceiling plan."),
    (4, 3, "HVAC indoor/outdoor routing coordination", "HVAC", "HVAC", "Coordinate AC piping, drain and outdoor unit routes."),
    (5, 1, "Partition layout marking and approval", "Carpentry", "Carpentry", "Mark partitions and confirm layout before execution."),
    (5, 2, "Electrical conduit work start", "Electrical", "Electrical", "Start conduit work after route approval."),
    (5, 3, "Fire line route verification", "Fire", "Fire", "Verify fire route and coordination requirements."),
]

CATEGORIES = [
    "Civil", "Electrical", "Carpentry", "Ceiling", "HVAC", "Fire", "IT", "Painting",
    "Flooring", "Glass", "Furniture", "Signage", "Housekeeping", "Quality", "Documentation", "Handover",
]


def generated_tasks():
    rows = list(DETAILED_TASKS)
    for day in range(6, 46):
        pivot = (day - 1) % len(CATEGORIES)
        rotation = CATEGORIES[pivot:] + CATEGORIES[:pivot]
        for offset, category in enumerate(rotation[:3]):
            vendor = category if category not in {"Quality", "Documentation", "Handover"} else "Internal"
            rows.append((
                day,
                offset + 1,
                f"{category} execution checkpoint day {day}.{offset + 1}",
                category,
                vendor,
                f"Complete the planned {category.lower()} activity for day {day} and report exceptions.",
            ))
    return rows


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            set_cell_width(cell, width)


def set_run_font(run, size=9, bold=False, color="172033"):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def configure_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size, color, before, after in [
        ("Heading 1", 16, "1D4ED8", 18, 10),
        ("Heading 2", 13, "1D4ED8", 14, 7),
        ("Heading 3", 12, "1E3A5F", 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)


def add_footer(section):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("Recovered reference | SiteOps V2 planning")
    set_run_font(run, size=8, color="64748B")


def add_title_block(doc):
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(3)
    run = kicker.add_run("SITEOPS | RECOVERED LEGACY REFERENCE")
    set_run_font(run, size=9, bold=True, color="1D4ED8")

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("45-Day Interior Fit-out Execution Template")
    set_run_font(run, size=25, bold=True, color="0F172A")

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(16)
    run = subtitle.add_run("Recovered from the initial Git commit for management and domain-expert review")
    set_run_font(run, size=12, color="475569")


def add_callout(doc):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF7ED")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Status: recovered legacy draft - not management-approved")
    set_run_font(r, size=10.5, bold=True, color="9A3412")
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(
        "Days 1-5 contain 15 specifically authored activities. Days 6-45 contain 120 generic category-rotation checkpoints generated by code. Use this document as source material only; validate and replace generic checkpoints before approving the V2 baseline."
    )
    set_run_font(r, size=9.5, color="7C2D12")


def add_metadata(doc):
    doc.add_heading("Recovery record", level=1)
    rows = [
        ("Git source", "Initial commit b1f5d6c2ded83dce9ecddef53d35690239d04ba5"),
        ("Original source file", "backend/app/seed.py"),
        ("Recovered schedule", "45 days, 3 activities per day, 135 activities total"),
        ("Authored content", "Days 1-5: 15 detailed activities"),
        ("Generated content", "Days 6-45: 120 generic rotating-category checkpoints"),
        ("Recommended use", "Reference input for the approved V2 45-day baseline workshop"),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    for label, value in rows:
        cells = table.add_row().cells
        cells[0].text = label
        cells[1].text = value
        set_cell_shading(cells[0], "E8EEF5")
        for run in cells[0].paragraphs[0].runs:
            set_run_font(run, size=9.5, bold=True, color="1E3A5F")
        for run in cells[1].paragraphs[0].runs:
            set_run_font(run, size=9.5)
    set_table_geometry(table, [2700, 6660])


def add_schedule(doc):
    doc.add_heading("Recovered day-by-day schedule", level=1)
    p = doc.add_paragraph(
        "The Origin column identifies whether an activity was explicitly authored in the legacy seed or generated by the category-rotation algorithm."
    )
    p.paragraph_format.space_after = Pt(8)

    headers = ["Day", "Activity", "Category", "Owner", "Origin"]
    widths = [560, 4550, 1320, 1250, 1680]
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, text in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = text
        set_cell_shading(cell, "E8EEF5")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for run in cell.paragraphs[0].runs:
            set_run_font(run, size=8.5, bold=True, color="1E3A5F")
    set_repeat_table_header(table.rows[0])

    for day, order, title, category, vendor, _note in generated_tasks():
        cells = table.add_row().cells
        values = [f"{day}", title, category, vendor, "Authored" if day <= 5 else "Generated"]
        for idx, value in enumerate(values):
            cells[idx].text = value
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if day <= 5:
                set_cell_shading(cells[idx], "F8FAFC")
            for paragraph in cells[idx].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run_font(run, size=8.2, bold=(idx == 0))
    set_table_geometry(table, widths)


def add_validation(doc):
    doc.add_heading("Required validation before V2 approval", level=1)
    items = [
        "Replace generic Days 6-45 checkpoints with site-operational tasks, dependencies and measurable acceptance criteria.",
        "Confirm whether every project uses one fixed baseline or whether approved project-type variants are permitted.",
        "Assign one Primary Responsible Employee to every active task and define optional supporting employees separately.",
        "Identify vendor or sub-vendor applicability without treating the vendor as the internal task owner.",
        "Define evidence, Supervisor verification and PM approval requirements for each task.",
        "Define reminders, escalation timing, absence reassignment and external approval dependencies.",
        "Obtain written management approval and version the baseline before implementation.",
    ]
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.add_run(item)


def build():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    configure_styles(doc)
    add_footer(section)
    add_title_block(doc)
    add_callout(doc)
    add_metadata(doc)
    doc.add_page_break()
    add_schedule(doc)
    add_validation(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT.resolve())


if __name__ == "__main__":
    build()
